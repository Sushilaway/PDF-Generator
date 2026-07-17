from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import json
import os
from datetime import datetime

def generate_docx(data, output_path):
    doc = Document()

    title = doc.add_heading(data.get('title', 'QA Test Report'), 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f"Project: {data.get('project', 'N/A')}")
    doc.add_paragraph(f"Tester: {data.get('tester', 'N/A')}")
    doc.add_paragraph(f"Date: {data.get('date', 'N/A')}")
    doc.add_paragraph()

    doc.add_heading('Summary', level=1)
    summary = data.get('summary', {})
    total = summary.get('total', 0)
    passed = summary.get('passed', 0)
    failed = summary.get('failed', 0)
    pass_rate = (passed / total * 100) if total > 0 else 0

    p = doc.add_paragraph()
    p.add_run(f'Total: {total}  |  ').bold = True
    p.add_run(f'\u2705 Passed: {passed}  |  ')
    p.add_run(f'\u274c Failed: {failed}  |  ')
    p.add_run(f'\U0001f4ca Pass Rate: {pass_rate:.1f}%')

    doc.add_paragraph()

    doc.add_heading('Test Cases', level=1)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'

    header_cells = table.rows[0].cells
    header_cells[0].text = 'ID'
    header_cells[1].text = 'Scenario'
    header_cells[2].text = 'Status'

    tests = data.get('tests', [])
    for test in tests:
        row = table.add_row().cells
        row[0].text = test.get('id', '')
        row[1].text = test.get('scenario', '')
        row[2].text = test.get('status', '')

        if test.get('status', '').lower() == 'pass':
            row[2].paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 128, 0)
        else:
            row[2].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 0, 0)

    doc.save(output_path)

def generate_docx_from_file(json_file='data.json', output_dir='output'):
    try:
        with open(json_file, 'r') as f:
            data = json.load(f)

        os.makedirs(output_dir, exist_ok=True)
        output_path = os.path.join(output_dir, 'report.docx')
        generate_docx(data, output_path)
        print(f"\u2705 DOCX generated: {output_path}")
        return output_path
    except Exception as e:
        print(f"\u274c Error generating DOCX: {e}")
        return None

if __name__ == "__main__":
    generate_docx_from_file()
